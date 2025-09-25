from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Dict, Any
import os
import json
import shutil
import requests
from urllib.parse import urlparse
from datetime import datetime
from bs4 import BeautifulSoup
import re
from dotenv import load_dotenv
from openai import OpenAI
from ai_functions import get_openai_functions, execute_function_call
from ai_control_functions import get_control_functions, execute_control_function

load_dotenv()

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

class DashboardQueryRequest(BaseModel):
    query: str

class UrlUploadRequest(BaseModel):
    url: str

class TextUploadRequest(BaseModel):
    filename: str
    content: str

class NoteInsertRequest(BaseModel):
    content: str

class ChatContinueRequest(BaseModel):
    message: str
    conversation_history: list = []
    dashboard_id: str = None

app = FastAPI(
    title="CB5 Capital Research Terminal API",
    description="Research Terminal Backend API for managing knowledge bases, generating AI-powered dashboards, and tracking query history.",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=os.getenv("CORS_ORIGINS"),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
async def root():
    """Health check endpoint."""
    return {"status": "ok", "message": "CB5 Capital Research Terminal API is running"}

@app.get("/health")
async def health():
    """Health check endpoint."""
    return {"status": "healthy", "version": "1.0.0"}

@app.get("/api/cases")
async def get_all_cases():
    """
    Get all available cases from ProjectLib directory.
    
    Returns:
        Dictionary containing list of cases with their metadata
    """
    project_lib_dir = "ProjectLib"
    
    # Check if ProjectLib directory exists
    if not os.path.exists(project_lib_dir):
        return {
            "cases": [],
            "total_count": 0
        }
    
    try:
        cases = []
        
        # List all JSON files in ProjectLib directory
        for filename in os.listdir(project_lib_dir):
            if filename.endswith('.json'):
                file_path = os.path.join(project_lib_dir, filename)
                
                try:
                    # Read the case file
                    with open(file_path, 'r', encoding='utf-8') as f:
                        case_data = json.load(f)
                    
                    # Extract case information
                    project_id = case_data.get('project_id', filename.replace('.json', ''))
                    project_name = case_data.get('project_name', 'Untitled Project')
                    
                    # Get file statistics
                    stat_info = os.stat(file_path)
                    
                    case_info = {
                        "project_id": project_id,
                        "project_name": project_name,
                        "filename": filename,
                        "created_at": stat_info.st_ctime,
                        "modified_at": stat_info.st_mtime
                    }
                    
                    cases.append(case_info)
                    
                except (json.JSONDecodeError, KeyError) as e:
                    print(f"Error reading case file {filename}: {e}")
                    continue
        
        # Sort cases by project_id
        cases.sort(key=lambda x: x.get('project_id', ''))
        
        return {
            "cases": cases,
            "total_count": len(cases)
        }
        
    except Exception as e:
        print(f"Error getting cases: {e}")
        raise HTTPException(status_code=500, detail=f"Error retrieving cases: {str(e)}")


@app.get("/api/cases/{case_name}/files")
async def list_case_files(case_name: str):
    case_dir = os.path.join("DataLib", case_name)
    if not os.path.exists(case_dir):
        return {"error": "Case not found"}
    
    files = []
    for filename in os.listdir(case_dir):
        file_path = os.path.join(case_dir, filename)
        if os.path.isfile(file_path):
            stat_info = os.stat(file_path)
            file_extension = filename.split('.')[-1].lower() if '.' in filename else 'unknown'
            
            files.append({
                "filename": filename,
                "file_type": file_extension,
                "size": stat_info.st_size,
                "modified_at": stat_info.st_mtime
            })
    
    return {"files": files}


@app.post("/api/cases/{case_name}/files/{filename}/generate-dashboard")
async def generate_dashboard_from_file(case_name: str, filename: str):
    """
    Generate dashboard items from a file using OpenAI function calling.
    
    Args:
        case_name: Name of the case (e.g., 'C1')
        filename: Name of the file to analyze
        
    Returns:
        Dictionary containing generated dashboard items
    """
    case_dir = os.path.join("DataLib", case_name)
    file_path = os.path.join(case_dir, filename)
    
    # Check if file exists
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    try:
        # Load research questions from ProjectLib case file
        research_questions = []
        project_file = os.path.join("ProjectLib", f"{case_name}.json")
        if os.path.exists(project_file):
            try:
                with open(project_file, 'r', encoding='utf-8') as f:
                    project_data = json.load(f)
                    research_questions = project_data.get('research_questions', [])
            except (json.JSONDecodeError, FileNotFoundError):
                print(f"Warning: Could not load research questions from {project_file}")
        
        # Read file content based on file type
        file_content = ""
        file_extension = filename.split('.')[-1].lower() if '.' in filename else 'unknown'
        
        if file_extension == 'json':
            # Handle JSON files
            with open(file_path, 'r', encoding='utf-8') as f:
                json_data = json.load(f)
                file_content = f"JSON Dataset: {filename}\nData Type: {type(json_data).__name__}\n\n"
                file_content += f"Content Summary:\n{json.dumps(json_data, indent=2)}"
        elif file_extension in ['docx', 'doc']:
            # Handle Word documents
            try:
                from docx import Document
                doc = Document(file_path)
                file_content = f"Word Document: {filename}\n\n"
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        file_content += paragraph.text + "\n"
            except ImportError:
                raise HTTPException(status_code=400, detail="Word document support not available. Please install python-docx: pip install python-docx")
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Unable to read Word document: {str(e)}")
        elif file_extension == 'pdf':
            # Handle PDF files
            try:
                import PyPDF2
                file_content = f"PDF Document: {filename}\n\n"
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page_num, page in enumerate(pdf_reader.pages):
                        text = page.extract_text()
                        if text.strip():
                            file_content += f"Page {page_num + 1}:\n{text}\n\n"
            except ImportError:
                raise HTTPException(status_code=400, detail="PDF support not available. Please install PyPDF2: pip install PyPDF2")
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Unable to read PDF document: {str(e)}")
        elif file_extension in ['txt', 'csv', 'md', 'py', 'js', 'html', 'xml']:
            # Handle text-based files
            with open(file_path, 'r', encoding='utf-8') as f:
                file_content = f.read()
        else:
            # Unsupported file type
            raise HTTPException(status_code=400, detail=f"Unsupported file type: .{file_extension}. Supported types: txt, csv, json, docx, pdf, md, py, js, html, xml")
        
        # Create the system prompt for the AI agent
        research_questions_text = ""
        if research_questions:
            research_questions_text = f"""
        
        RESEARCH QUESTIONS TO ADDRESS:
        The goal of this analysis is to provide data and insights that help answer these research questions:
        {chr(10).join(f"- {question}" for question in research_questions)}
        
        When creating dashboard components, prioritize information that directly addresses or provides data relevant to answering these research questions. Focus on extracting insights, metrics, and analysis that would help stakeholders understand and respond to these key research areas.
        """

        system_prompt = f"""You are an AI agent specialized in creating comprehensive dashboard components from business documents.

        Your task is to analyze the provided file content and create multiple relevant dashboard items that provide valuable insights for business analysis.{research_questions_text}

        Guidelines:
        1. Create multiple dashboard components (aim for 4-8 components)
        2. Use appropriate component types for different types of data:
        - Use metric_card for key numbers, KPIs, and important values
        - Use data_table for structured data that can be organized in rows/columns
        - Use financial_chart for numerical data that can be visualized
        - Use list_items for bullet points, key factors, or lists
        - Use text_analysis for comprehensive analysis with insights
        - Use competitor_analysis for competitor-related information
        - Use risk_assessment for risk-related information
        - Use short_text for brief summaries
        - Use long_text for detailed explanations
        - Use progress_bar for completion rates or percentages

        3. Extract specific numbers, percentages, and data points
        4. Identify key insights, trends, and important information
        5. Create components that would be valuable for business decision-making
        6. Always include the source filename in your function calls
        7. Provide meaningful titles and clear, concise content
        8. Vary the component sizes appropriately (small for metrics, medium/large for detailed components)
        9. PRIORITIZE information that helps answer the research questions listed above

        Remember to call the appropriate functions to create each dashboard component. Be thorough and create a comprehensive dashboard that covers all important aspects of the file content, especially focusing on data that addresses the research questions.
        """

        # Create the user prompt with file content
        research_context = ""
        if research_questions:
            research_context = f"\n\nRemember to focus on information that helps answer these research questions:\n{chr(10).join(f'- {question}' for question in research_questions)}"
        
        user_prompt = f"""Please analyze the following file content and create comprehensive dashboard components:

File: {filename}
Case: {case_name}
Content: {file_content}{research_context}

Please create multiple relevant dashboard components that provide valuable business insights from this content, especially focusing on data that addresses the research questions. Use the available functions to create each component."""

        # Make the OpenAI API call with function calling
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            tools=get_openai_functions(),
            tool_choice="auto",
            temperature=0.3,
            max_tokens=4000
        )
        
        # Process function calls and generate individual dashboard items
        dashboard_items = []
        
        if response.choices[0].message.tool_calls:
            for tool_call in response.choices[0].message.tool_calls:
                function_name = tool_call.function.name
                function_args = json.loads(tool_call.function.arguments)
                
                # Execute the function call to create the dashboard component
                try:
                    component = execute_function_call(function_name, function_args)
                    
                    # Create individual dashboard item for each component
                    item = {
                        "id": f"{case_name}_{filename}_{function_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                        "source_file": filename,
                        "created_at": datetime.now().isoformat(),
                        "analysis_type": "file_analysis",
                        "component": component,
                        "metadata": {
                            "file_type": file_extension,
                            "file_size": os.path.getsize(file_path),
                            "component_type": component["type"],
                            "analysis_query": f"Extract {component['type']} insights from {filename}"
                        }
                    }
                    dashboard_items.append(item)
                    
                except Exception as e:
                    print(f"Error executing function {function_name}: {e}")
                    continue
        
        # Save to items.json file (individual dashboard items for this case)
        dashboard_case_dir = os.path.join("DashboardLib", case_name)
        
        # Create DashboardLib/CaseName directory if it doesn't exist
        os.makedirs(dashboard_case_dir, exist_ok=True)
        
        items_file = os.path.join(dashboard_case_dir, "items.json")
        
        # Load existing items or create new list
        existing_items = []
        if os.path.exists(items_file):
            try:
                with open(items_file, 'r', encoding='utf-8') as f:
                    existing_items = json.load(f)
            except (json.JSONDecodeError, FileNotFoundError):
                existing_items = []
        
        # Add new dashboard items
        existing_items.extend(dashboard_items)
        
        # Save updated items
        with open(items_file, 'w', encoding='utf-8') as f:
            json.dump(existing_items, f, indent=2, ensure_ascii=False)
        
        # Run optimization after generation
        optimization_result = None
        try:
            print(f"Running optimization for case {case_name}...")
            optimization_result = await run_dashboard_optimization(case_name)
            print(f"Optimization completed: {optimization_result.get('message', 'Success')}")
        except Exception as e:
            print(f"Optimization failed but generation succeeded: {e}")
            # Don't fail the entire request if optimization fails
        
        return {
            "success": True,
            "items_created": len(dashboard_items),
            "items": dashboard_items,
            "optimization": optimization_result,
            "message": f"Successfully generated {len(dashboard_items)} dashboard items from {filename}" + 
                      (f" and optimized to {optimization_result.get('final_item_count', 'unknown')} items" if optimization_result and optimization_result.get('success') else "")
        }
        
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="File not found")
    except UnicodeDecodeError:
        raise HTTPException(status_code=400, detail="Unable to read file content - file may be binary or corrupted")
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid JSON file format")
    except Exception as e:
        print(f"Error generating dashboard: {e}")
        raise HTTPException(status_code=500, detail=f"Error generating dashboard: {str(e)}")


async def run_dashboard_optimization(case_name: str) -> Dict[str, Any]:
    """
    Internal function to run dashboard optimization for a case.
    
    Args:
        case_name: Name of the case to optimize
        
    Returns:
        Dictionary containing optimization results
    """
    # Check if dashboard items exist
    dashboard_case_dir = os.path.join("DashboardLib", case_name)
    items_file = os.path.join(dashboard_case_dir, "items.json")
    
    if not os.path.exists(items_file):
        return {
            "success": False,
            "error": f"No dashboard items found for case {case_name}"
        }
    
    try:
        # Load existing dashboard items to analyze
        with open(items_file, 'r', encoding='utf-8') as f:
            items = json.load(f)
        
        if not items:
            return {
                "success": True,
                "message": "No items to optimize",
                "optimization_actions": [],
                "final_item_count": 0
            }
        
        # Pre-processing: Remove obvious duplicates using rule-based approach
        print(f"Starting with {len(items)} items")
        deduplicated_items = []
        seen_combinations = set()
        duplicates_removed = 0
        
        for item in items:
            component = item.get('component', {})
            # Create a unique key based on type, title/label, and value
            item_type = component.get('type', '')
            title = component.get('title', '') or component.get('label', '')
            value = str(component.get('value', ''))
            
            # For metric cards, use type + label + value as key
            if item_type == 'metric_card':
                unique_key = f"{item_type}|{title}|{value}"
            # For other types, use type + title as key
            else:
                unique_key = f"{item_type}|{title}"
            
            if unique_key not in seen_combinations:
                seen_combinations.add(unique_key)
                deduplicated_items.append(item)
            else:
                duplicates_removed += 1
                print(f"Removed duplicate: {item_type} - {title}")
        
        # Save deduplicated items back to file
        if duplicates_removed > 0:
            with open(items_file, 'w', encoding='utf-8') as f:
                json.dump(deduplicated_items, f, indent=2, ensure_ascii=False)
            print(f"Pre-processing removed {duplicates_removed} obvious duplicates")
            items = deduplicated_items
        
        # Create the system prompt for the control agent
        system_prompt = """You are an AI control agent that MUST aggressively optimize dashboard items by removing duplicates and redundancy.

        CRITICAL: You MUST take action every iteration. Do not just list items - you must actively delete, update, or consolidate items.

        Your mandatory workflow:
        1. FIRST: Call list_dashboard_items() to see all items
        2. IMMEDIATELY identify obvious duplicates (same type, same label/title, same values)
        3. AGGRESSIVELY delete duplicate items using delete_dashboard_item()
        4. Look for similar items that can be merged using create_consolidated_item()
        5. Continue until no more optimization is possible
        6. FINALLY: Call mark_optimization_complete()

        RULES FOR DELETION:
        - If multiple metric_card items have the same label (like "Global Market Size 2025"), DELETE all but the most recent one
        - If multiple items show essentially the same information, DELETE duplicates immediately
        - If items have identical titles and types, DELETE the older ones
        - Be very aggressive - reduce the total number of items significantly

        You MUST make deletion or consolidation actions every single iteration. Do not hesitate to delete items.

        Current target: Reduce the dashboard items by at least 50% by removing obvious duplicates and redundant information."""

        # Create the user prompt with case information
        user_prompt = f"""Please optimize the dashboard items for case {case_name}.

        Current item count: {len(items)}

        Your goal is to create an efficient, well-organized set of dashboard items by:
        1. Removing duplicates and redundant items
        2. Merging similar items into consolidated components
        3. Ensuring all valuable information is preserved
        4. Creating a clean, coherent dashboard structure

        Start by listing all items to understand the current state, then work systematically to optimize the collection."""

        # Initialize optimization tracking
        optimization_actions = []
        max_iterations = 5  # Reduced for integration workflow
        iteration_count = 0
        optimization_complete = False
        
        while not optimization_complete and iteration_count < max_iterations:
            iteration_count += 1
            print(f"Control agent iteration {iteration_count}")
            
            # Make the OpenAI API call with function calling
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt if iteration_count == 1 else "Continue optimizing the dashboard items. If no more optimization is needed, mark the process as complete."}
                ],
                tools=get_control_functions(),
                tool_choice="auto",
                temperature=0.2,
                max_tokens=2000
            )
            
            # Process function calls
            iteration_actions = []
            
            if response.choices[0].message.tool_calls:
                for tool_call in response.choices[0].message.tool_calls:
                    function_name = tool_call.function.name
                    function_args = json.loads(tool_call.function.arguments)
                    
                    # Execute the function call
                    try:
                        print(f"Executing {function_name} with args: {function_args}")
                        result = execute_control_function(function_name, function_args)
                        print(f"Function {function_name} result: {result}")
                        
                        # Track the action
                        action = {
                            "iteration": iteration_count,
                            "function": function_name,
                            "arguments": function_args,
                            "result": result
                        }
                        iteration_actions.append(action)
                        optimization_actions.append(action)
                        
                        # Check if optimization is complete
                        if function_name == "mark_optimization_complete" and result.get("success"):
                            optimization_complete = True
                            break
                            
                    except Exception as e:
                        print(f"Error executing control function {function_name}: {e}")
                        continue
            else:
                # No function calls - agent might be done or stuck
                break
            
            # If no actions were taken this iteration, we're likely done
            if not iteration_actions:
                break
        
        # Get final item count
        final_items = []
        if os.path.exists(items_file):
            try:
                with open(items_file, 'r', encoding='utf-8') as f:
                    final_items = json.load(f)
            except:
                pass
        
        return {
            "success": True,
            "case_name": case_name,
            "original_item_count": len(items) + duplicates_removed,  # Include pre-processing removals
            "final_item_count": len(final_items),
            "items_removed": (len(items) + duplicates_removed) - len(final_items),
            "duplicates_removed_preprocessing": duplicates_removed,
            "iterations": iteration_count,
            "optimization_actions": optimization_actions,
            "optimization_completed": optimization_complete,
            "message": f"Optimization completed. Reduced from {len(items) + duplicates_removed} to {len(final_items)} items ({duplicates_removed} by pre-processing, {len(items) - len(final_items)} by AI) in {iteration_count} iterations."
        }
        
    except Exception as e:
        print(f"Error during optimization: {e}")
        return {
            "success": False,
            "error": f"Error optimizing dashboard items: {str(e)}"
        }


@app.post("/api/cases/{case_name}/generate-dashboard")
async def generate_dashboard_from_query(case_name: str, request: DashboardQueryRequest):
    """
    Generate a dashboard from a query using existing dashboard items.
    
    Args:
        case_name: Name of the case (e.g., 'C1')
        request: Request containing the user query
        
    Returns:
        Dictionary containing the generated dashboard configuration
    """
    query = request.query
    # Path to the dashboard items for this case
    dashboard_case_dir = os.path.join("DashboardLib", case_name)
    items_file = os.path.join(dashboard_case_dir, "items.json")
    
    # Check if items.json exists
    if not os.path.exists(items_file):
        raise HTTPException(status_code=404, detail=f"No dashboard items found for case {case_name}")
    
    try:
        # Load research questions from ProjectLib case file
        research_questions = []
        project_file = os.path.join("ProjectLib", f"{case_name}.json")
        if os.path.exists(project_file):
            try:
                with open(project_file, 'r', encoding='utf-8') as f:
                    project_data = json.load(f)
                    research_questions = project_data.get('research_questions', [])
            except (json.JSONDecodeError, FileNotFoundError):
                print(f"Warning: Could not load research questions from {project_file}")
        
        # Load existing dashboard items
        with open(items_file, 'r', encoding='utf-8') as f:
            available_items = json.load(f)
        
        if not available_items:
            raise HTTPException(status_code=404, detail=f"No dashboard items available for case {case_name}")
        
        # Create a summary of available items for the AI
        items_summary = []
        for item in available_items:
            component = item.get('component', {})
            summary = {
                "id": item.get('id'),
                "type": component.get('type'),
                "title": component.get('title', component.get('label', 'Untitled')),
                "source_file": item.get('source_file'),
                "created_at": item.get('created_at')
            }
            items_summary.append(summary)
        
        # Create the system prompt for dashboard generation
        research_questions_context = ""
        if research_questions:
            research_questions_context = f"""
            
            RESEARCH QUESTIONS CONTEXT:
            This case is focused on answering these research questions:
            {chr(10).join(f"            - {question}" for question in research_questions)}
            
            When selecting dashboard items, prioritize those that provide data and insights relevant to these research questions. Consider how each item contributes to answering these key research areas.
            """

        system_prompt = f"""You are an AI agent specialized in creating dashboard configurations from existing dashboard items.

            Your task is to:
            1. Analyze the user's query
            2. Select the most relevant dashboard items from the available items
            3. Create a dashboard layout that best addresses the user's query
            4. Return a JSON dashboard configuration{research_questions_context}

            Guidelines:
            - Select 3-8 dashboard items that are most relevant to the query
            - Arrange items in a logical layout (2-4 columns work best)
            - Prioritize items that directly answer or relate to the user's query and research questions
            - Consider the visual balance - mix different component types
            - Use appropriate sizing (small for metrics, medium/large for detailed components)

            Available dashboard item types:
            - metric_card: KPIs and key numbers
            - data_table: Structured data
            - financial_chart: Charts and graphs
            - list_items: Bulleted lists
            - text_analysis: Analysis content
            - competitor_analysis: Competitor comparisons
            - risk_assessment: Risk analysis
            - short_text: Brief summaries
            - long_text: Detailed content
            - progress_bar: Progress indicators

            You must respond with a JSON object in this exact format:
            {
            "title": "Dashboard Title Based on Query",
            "subtitle": "Brief description of what this dashboard shows",
            "layout": "grid",
            "columns": 3,
            "components": [
                {
                "item_id": "id_of_selected_item",
                "size": "small|medium|large",
                "position": {
                    "row": 1,
                    "col": 1
                }
                }
            ]
            }
        """

        # Create the user prompt with available items and query
        research_prompt_context = ""
        if research_questions:
            research_prompt_context = f"""
            
            Research Questions for this case:
            {chr(10).join(f'            - {question}' for question in research_questions)}
            
            Consider how your dashboard selection can provide insights to answer these research questions."""

        user_prompt = f"""Please create a dashboard configuration for the following query:

            Query: "{query}"
            Case: {case_name}{research_prompt_context}

            Available Dashboard Items:
            {json.dumps(items_summary, indent=2)}

        Please analyze the query and select the most relevant items to create a comprehensive dashboard that addresses the user's needs and helps answer the research questions. Return only the JSON configuration."""

        # Make the OpenAI API call
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3,
            max_tokens=2000,
            response_format={"type": "json_object"}
        )
        
        # Parse the dashboard configuration
        dashboard_config = json.loads(response.choices[0].message.content)
        
        # Validate and enrich the dashboard configuration with actual component data
        enriched_components = []
        selected_item_ids = [comp.get('item_id') for comp in dashboard_config.get('components', [])]
        
        for comp_config in dashboard_config.get('components', []):
            item_id = comp_config.get('item_id')
            
            # Find the actual item
            actual_item = next((item for item in available_items if item.get('id') == item_id), None)
            if actual_item:
                # Merge the AI's layout decisions with the actual component
                enriched_component = actual_item['component'].copy()
                enriched_component['size'] = comp_config.get('size', enriched_component.get('size', 'medium'))
                
                # Add position if provided
                if 'position' in comp_config:
                    enriched_component['position'] = comp_config['position']
                
                enriched_components.append(enriched_component)
        
        # Generate unique dashboard filename first
        dashboard_id = f"dashboard_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        # Create the final dashboard
        final_dashboard = {
            "title": dashboard_config.get('title', f'Dashboard for: {query}'),
            "subtitle": dashboard_config.get('subtitle', f'Case {case_name}'),
            "layout": dashboard_config.get('layout', 'grid'),
            "columns": dashboard_config.get('columns', 3),
            "components": enriched_components,
            "conversation_history": [],  # Initialize empty conversation history
            "metadata": {
                "query": query,
                "case_name": case_name,
                "research_questions": research_questions,
                "items_selected": len(enriched_components),
                "total_items_available": len(available_items),
                "selected_item_ids": selected_item_ids,
                "created_at": datetime.now().isoformat(),
                "dashboard_id": dashboard_id  # Add dashboard_id to metadata
            }
        }
        
        # Save the dashboard to QueryLib/CaseName
        query_case_dir = os.path.join("QueryLib", case_name)
        os.makedirs(query_case_dir, exist_ok=True)
        dashboard_file = os.path.join(query_case_dir, f"{dashboard_id}.json")
        
        # Save the dashboard
        with open(dashboard_file, 'w', encoding='utf-8') as f:
            json.dump(final_dashboard, f, indent=2, ensure_ascii=False)
        
        return {
            "success": True,
            "dashboard": final_dashboard,
            "dashboard_id": dashboard_id,
            "message": f"Successfully generated dashboard with {len(enriched_components)} components"
        }
        
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=500, detail=f"Error parsing dashboard items or AI response: {str(e)}")
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail=f"Dashboard items file not found for case {case_name}")
    except Exception as e:
        print(f"Error generating dashboard from query: {e}")
        raise HTTPException(status_code=500, detail=f"Error generating dashboard: {str(e)}")


@app.get("/api/cases/{case_name}/queries")
async def get_case_queries(case_name: str):
    """
    Get all queries/dashboards for a given case from QueryLib.
    
    Args:
        case_name: Name of the case (e.g., 'C1')
        
    Returns:
        Dictionary containing list of queries with metadata
    """
    # Path to the queries for this case
    query_case_dir = os.path.join("QueryLib", case_name)
    
    # Check if QueryLib directory exists for this case
    if not os.path.exists(query_case_dir):
        return {
            "case_name": case_name,
            "queries": [],
            "total_count": 0
        }
    
    try:
        queries = []
        
        # List all dashboard JSON files in the directory
        for filename in os.listdir(query_case_dir):
            if filename.endswith('.json') and filename.startswith('dashboard_'):
                file_path = os.path.join(query_case_dir, filename)
                
                try:
                    # Read the dashboard file
                    with open(file_path, 'r', encoding='utf-8') as f:
                        dashboard_data = json.load(f)
                    
                    # Extract query information
                    metadata = dashboard_data.get('metadata', {})
                    dashboard_id = filename.replace('.json', '')
                    
                    query_info = {
                        "dashboard_id": dashboard_id,
                        "query": metadata.get('query', dashboard_data.get('title', 'Untitled Query')),
                        "title": dashboard_data.get('title', 'Untitled Dashboard'),
                        "subtitle": dashboard_data.get('subtitle', ''),
                        "created_at": metadata.get('created_at', ''),
                        "items_selected": metadata.get('items_selected', 0),
                        "total_items_available": metadata.get('total_items_available', 0),
                        "components_count": len(dashboard_data.get('components', [])),
                        "file_path": filename
                    }
                    
                    queries.append(query_info)
                    
                except (json.JSONDecodeError, KeyError) as e:
                    print(f"Error reading dashboard file {filename}: {e}")
                    continue
        
        # Sort queries by creation date (newest first)
        queries.sort(key=lambda x: x.get('created_at', ''), reverse=True)
        
        return {
            "case_name": case_name,
            "queries": queries,
            "total_count": len(queries)
        }
        
    except Exception as e:
        print(f"Error getting case queries: {e}")
        raise HTTPException(status_code=500, detail=f"Error retrieving queries: {str(e)}")


@app.get("/api/settings")
async def get_settings():
    """
    Get application settings from settings.json.
    
    Returns:
        Dictionary containing application settings
    """
    settings_file = "settings.json"
    
    try:
        if os.path.exists(settings_file):
            with open(settings_file, 'r', encoding='utf-8') as f:
                settings = json.load(f)
            return {
                "success": True,
                "settings": settings
            }
        else:
            # Return default settings if file doesn't exist
            default_settings = {
                "current_project": None,
                "theme": "dark"
            }
            return {
                "success": True,
                "settings": default_settings
            }
    except Exception as e:
        print(f"Error reading settings: {e}")
        raise HTTPException(status_code=500, detail=f"Error reading settings: {str(e)}")


@app.post("/api/settings")
async def update_settings(settings: Dict[str, Any]):
    """
    Update application settings in settings.json.
    
    Args:
        settings: Dictionary containing settings to update
        
    Returns:
        Dictionary containing success status
    """
    settings_file = "settings.json"
    
    try:
        # Read existing settings or create default
        existing_settings = {}
        if os.path.exists(settings_file):
            with open(settings_file, 'r', encoding='utf-8') as f:
                existing_settings = json.load(f)
        
        # Update with new settings
        existing_settings.update(settings)
        
        # Write back to file
        with open(settings_file, 'w', encoding='utf-8') as f:
            json.dump(existing_settings, f, indent=2, ensure_ascii=False)
        
        return {
            "success": True,
            "message": "Settings updated successfully",
            "settings": existing_settings
        }
        
    except Exception as e:
        print(f"Error updating settings: {e}")
        raise HTTPException(status_code=500, detail=f"Error updating settings: {str(e)}")


@app.get("/api/cases/{case_name}/queries/{dashboard_id}")
async def get_dashboard_by_id(case_name: str, dashboard_id: str):
    """
    Get a specific dashboard by its ID from QueryLib.
    
    Args:
        case_name: Name of the case (e.g., 'C1')
        dashboard_id: ID of the dashboard (e.g., 'dashboard_20250923_112850')
        
    Returns:
        Dictionary containing the dashboard data
    """
    # Path to the specific dashboard file
    query_case_dir = os.path.join("QueryLib", case_name)
    dashboard_file = os.path.join(query_case_dir, f"{dashboard_id}.json")
    
    # Check if dashboard file exists
    if not os.path.exists(dashboard_file):
        raise HTTPException(status_code=404, detail=f"Dashboard {dashboard_id} not found for case {case_name}")
    
    try:
        # Read the dashboard file
        with open(dashboard_file, 'r', encoding='utf-8') as f:
            dashboard_data = json.load(f)
        
        return {
            "success": True,
            "dashboard": dashboard_data,
            "dashboard_id": dashboard_id
        }
        
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=500, detail=f"Error parsing dashboard file: {str(e)}")
    except Exception as e:
        print(f"Error loading dashboard {dashboard_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Error loading dashboard: {str(e)}")


@app.get("/api/cases/{case_name}/sources/{filename}")
async def get_source_content(case_name: str, filename: str):
    """
    Get the full content of a source file referenced in dashboard components.
    
    Args:
        case_name: Name of the case (e.g., 'C1')
        filename: Name of the source file
        
    Returns:
        Dictionary containing the full file content
    """
    # Path to the case directory - check both DataLib and ProjectLib
    case_dir = os.path.join("DataLib", case_name)
    if not os.path.exists(case_dir):
        case_dir = os.path.join("ProjectLib", case_name)
    
    # Check if case directory exists
    if not os.path.exists(case_dir):
        raise HTTPException(status_code=404, detail=f"Case {case_name} not found")
    
    # Look for the file in the case directory
    file_path = os.path.join(case_dir, filename)
    
    # Check if file exists
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail=f"Source file {filename} not found in case {case_name}")
    
    try:
        # Get file stats first
        stat = os.stat(file_path)
        file_size = stat.st_size
        modified_time = stat.st_mtime
        
        # Determine file type and read accordingly
        file_extension = filename.lower().split('.')[-1] if '.' in filename else ''
        
        if file_extension in ['docx', 'doc']:
            # Handle Word documents
            try:
                from docx import Document
                doc = Document(file_path)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
                file_type = "word_document"
            except ImportError:
                # If python-docx is not installed, try alternative approach
                raise HTTPException(status_code=400, detail="Cannot read Word documents - python-docx library not available. Please convert to .txt format.")
            except Exception as docx_error:
                raise HTTPException(status_code=400, detail=f"Error reading Word document: {str(docx_error)}")
                
        elif file_extension in ['pdf']:
            raise HTTPException(status_code=400, detail="PDF files not supported yet. Please convert to .txt format.")
            
        elif file_extension in ['txt', 'csv', 'json', 'py', 'js', 'html', 'css', 'md', 'rst']:
            # Handle text files
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                file_type = "text_file"
            except UnicodeDecodeError:
                # Try with different encodings
                for encoding in ['latin-1', 'cp1252']:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            content = f.read()
                        file_type = "text_file"
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    raise HTTPException(status_code=400, detail="File encoding not supported")
        else:
            # Unknown file type - try to read as text first
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                file_type = "unknown_text"
            except UnicodeDecodeError:
                raise HTTPException(status_code=400, detail=f"Unsupported file type: .{file_extension}")
        
        return {
            "success": True,
            "filename": filename,
            "content": content,
            "file_size": file_size,
            "modified_time": modified_time,
            "case_name": case_name,
            "file_type": file_type
        }
        
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        print(f"Error reading source file {filename}: {e}")
        raise HTTPException(status_code=500, detail=f"Error reading source file: {str(e)}")

@app.post("/api/cases/{case_name}/upload/file")
async def upload_file(case_name: str, file: UploadFile = File(...)):
    """
    Upload a file to the case's DataLib directory.
    """
    try:
        # Ensure DataLib directory exists
        data_lib_dir = os.path.join("DataLib", case_name)
        os.makedirs(data_lib_dir, exist_ok=True)
        
        # Save the uploaded file
        file_path = os.path.join(data_lib_dir, file.filename)
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        return {
            "success": True,
            "message": f"File {file.filename} uploaded successfully",
            "filename": file.filename,
            "case_name": case_name,
            "file_size": os.path.getsize(file_path)
        }
        
    except Exception as e:
        print(f"Error uploading file: {e}")
        raise HTTPException(status_code=500, detail=f"Error uploading file: {str(e)}")

@app.post("/api/cases/{case_name}/upload/url")
async def upload_url(case_name: str, request: UrlUploadRequest):
    """
    Scrape content from a webpage URL and save it as a text file.
    """
    try:
        # Ensure DataLib directory exists
        data_lib_dir = os.path.join("DataLib", case_name)
        os.makedirs(data_lib_dir, exist_ok=True)
        
        # Download the webpage content
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(request.url, headers=headers, timeout=30)
        response.raise_for_status()
        
        # Check content type
        content_type = response.headers.get('content-type', '').lower()
        
        if 'text/html' not in content_type and 'text/plain' not in content_type:
            # Handle non-HTML content (PDFs, documents, etc.)
            text_content = f"Binary/Document Content from {request.url}\nContent-Type: {content_type}\nFile size: {len(response.content)} bytes\n\nNote: This appears to be a binary file or non-HTML document."
            title_text = f"Document from {urlparse(request.url).netloc}"
        else:
            # Parse HTML content
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove unwanted elements
            for element in soup(["script", "style", "nav", "footer", "header", "aside", "iframe", "noscript"]):
                element.decompose()
                
            # Remove elements with common ad/navigation classes
            unwanted_classes = ['advertisement', 'ads', 'sidebar', 'navigation', 'menu', 'social', 'share']
            for class_name in unwanted_classes:
                for element in soup.find_all(attrs={'class': lambda x: x and any(unwanted in ' '.join(x).lower() for unwanted in unwanted_classes)}):
                    element.decompose()
            
            # Extract title
            title = soup.find('title')
            title_text = title.get_text().strip() if title else f"Page from {urlparse(request.url).netloc}"
            
            # Extract main content using multiple strategies
            main_content = None
            
            # Strategy 1: Look for semantic HTML5 elements and common content classes
            content_selectors = [
                'main', 'article', '[role="main"]',
                '.content', '#content', '.main-content', '.page-content',
                '.post-content', '.entry-content', '.article-content', 
                '.story-body', '.article-body', '.post-body'
            ]
            
            for selector in content_selectors:
                elements = soup.select(selector)
                if elements:
                    # Get the largest content block
                    main_content = max(elements, key=lambda x: len(x.get_text()))
                    break
            
            # Strategy 2: If no main content found, look for the largest text block
            if not main_content:
                all_divs = soup.find_all(['div', 'section', 'p'])
                if all_divs:
                    main_content = max(all_divs, key=lambda x: len(x.get_text()))
            
            # Strategy 3: Fall back to body
            if not main_content:
                main_content = soup.find('body') or soup
            
            # Extract and clean text
            text_content = main_content.get_text(separator='\n', strip=True)
            
            # Clean up the text more aggressively
            lines = text_content.split('\n')
            cleaned_lines = []
            
            for line in lines:
                line = line.strip()
                # Skip very short lines (likely navigation/ads)
                if len(line) > 3:
                    # Skip lines that look like navigation
                    if not any(nav_word in line.lower() for nav_word in ['home', 'menu', 'login', 'search', 'subscribe', 'follow us']):
                        cleaned_lines.append(line)
            
            text_content = '\n'.join(cleaned_lines)
            
            # Final cleanup
            text_content = re.sub(r'\n\s*\n\s*\n+', '\n\n', text_content)  # Max 2 consecutive newlines
            text_content = text_content.strip()
        
        # Create content with metadata
        scraped_content = f"""Title: {title_text}
URL: {request.url}
Scraped on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

Content:
{text_content}
"""
        
        # Generate filename from title or URL
        safe_title = re.sub(r'[^\w\s-]', '', title_text)
        safe_title = re.sub(r'[\s_-]+', '_', safe_title)
        safe_title = safe_title[:50]  # Limit length
        
        if not safe_title:
            parsed_url = urlparse(request.url)
            safe_title = parsed_url.netloc.replace('.', '_')
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"url_{safe_title}_{timestamp}.txt"
        
        # Save the scraped content
        file_path = os.path.join(data_lib_dir, filename)
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(scraped_content)
        
        return {
            "success": True,
            "message": f"Webpage content scraped and saved as {filename}",
            "filename": filename,
            "case_name": case_name,
            "url": request.url,
            "title": title_text,
            "content_type": content_type,
            "file_size": os.path.getsize(file_path),
            "content_length": len(text_content)
        }
        
    except requests.RequestException as e:
        print(f"Error fetching URL {request.url}: {e}")
        raise HTTPException(status_code=400, detail=f"Error fetching URL: {str(e)}")
    except Exception as e:
        print(f"Error scraping URL content: {e}")
        raise HTTPException(status_code=500, detail=f"Error scraping URL content: {str(e)}")

@app.post("/api/cases/{case_name}/upload/text")
async def upload_text(case_name: str, request: TextUploadRequest):
    """
    Create a text file with the provided content in the case's DataLib directory.
    """
    try:
        # Ensure DataLib directory exists
        data_lib_dir = os.path.join("DataLib", case_name)
        os.makedirs(data_lib_dir, exist_ok=True)
        
        # Add .txt extension if not present
        filename = request.filename
        if not filename.endswith('.txt'):
            filename += '.txt'
        
        # Save the text content
        file_path = os.path.join(data_lib_dir, filename)
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(request.content)
        
        return {
            "success": True,
            "message": f"Text file {filename} created successfully",
            "filename": filename,
            "case_name": case_name,
            "file_size": os.path.getsize(file_path)
        }
        
    except Exception as e:
        print(f"Error creating text file: {e}")
        raise HTTPException(status_code=500, detail=f"Error creating text file: {str(e)}")

@app.post("/api/cases/{case_name}/upload/note")
async def insert_note(case_name: str, request: NoteInsertRequest):
    """
    Insert a timestamped note into the case's DataLib directory.
    """
    try:
        # Ensure DataLib directory exists
        data_lib_dir = os.path.join("DataLib", case_name)
        os.makedirs(data_lib_dir, exist_ok=True)
        
        # Generate timestamped filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"note_{timestamp}.txt"
        
        # Save the note content
        file_path = os.path.join(data_lib_dir, filename)
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(request.content)
        
        return {
            "success": True,
            "message": f"Note {filename} inserted successfully",
            "filename": filename,
            "case_name": case_name,
            "timestamp": timestamp,
            "file_size": os.path.getsize(file_path)
        }
        
    except Exception as e:
        print(f"Error inserting note: {e}")
        raise HTTPException(status_code=500, detail=f"Error inserting note: {str(e)}")

@app.post("/api/cases/{case_name}/chat/continue")
async def continue_chat(case_name: str, request: ChatContinueRequest):
    """
    Continue a conversation about the current dashboard data without generating new dashboards.
    This is for follow-up questions and clarifications.
    """
    try:
        # Get the case data directory  
        data_lib_dir = os.path.join("DataLib", case_name)
        if not os.path.exists(data_lib_dir):
            raise HTTPException(status_code=404, detail=f"Case {case_name} not found")
        
        # Get available files for context
        available_files = []
        for filename in os.listdir(data_lib_dir):
            if filename.endswith(('.txt', '.docx', '.pdf', '.csv')):
                available_files.append(filename)
        
        # Build conversation context
        conversation_context = ""
        if request.conversation_history:
            conversation_context = "\n".join([
                f"{'User' if msg.get('type') == 'user' else 'Assistant'}: {msg.get('content', '')}"
                for msg in request.conversation_history[-6:]  # Last 6 messages for context
            ])
        
        # Read actual content from available files for analysis
        file_contents = {}
        for filename in available_files[:5]:  # Limit to first 5 files to avoid token limits
            try:
                file_path = os.path.join(data_lib_dir, filename)
                if filename.endswith('.txt'):
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()[:3000]  # First 3000 characters
                        file_contents[filename] = content
            except Exception as e:
                print(f"Error reading file {filename}: {e}")
                continue

        # Create the enhanced chat prompt with actual data
        files_data = ""
        if file_contents:
            for filename, content in file_contents.items():
                files_data += f"\n\n=== {filename} ===\n{content}\n"

        chat_prompt = f"""You are an expert business analyst with access to real data about case {case_name}.

ACTUAL DATA AVAILABLE:
{files_data if files_data else "No readable text files found."}

CONVERSATION CONTEXT:
{conversation_context}

USER QUESTION: {request.message}

INSTRUCTIONS:
- You are an expert analyst, not a pointer or guide
- Analyze the actual data above to answer the user's question directly
- Provide specific numbers, trends, insights from the data
- If you can't find specific information in the data, say so clearly
- Cite the exact filename when referencing specific data points
- Give concrete analysis, not suggestions about where to look
- Be direct and authoritative in your analysis

Analyze the data and provide a comprehensive answer to the user's question."""

        # Call OpenAI for the chat response
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {
                    "role": "system", 
                    "content": "You are an expert business analyst. Analyze data directly and provide specific insights with numbers and concrete findings. Never just point to where information might be - always analyze and conclude."
                },
                {
                    "role": "user", 
                    "content": chat_prompt
                }
            ],
            max_tokens=800,
            temperature=0.7
        )
        
        ai_response = response.choices[0].message.content
        
        # Extract potential source references (basic implementation)
        sources = []
        for filename in available_files:
            if filename.lower().replace('.txt', '').replace('.docx', '').replace('.pdf', '') in ai_response.lower():
                sources.append(filename)
        
        # Save conversation history to dashboard if dashboard_id is provided
        if request.dashboard_id:
            try:
                # Load existing dashboard
                query_case_dir = os.path.join("QueryLib", case_name)
                dashboard_file = os.path.join(query_case_dir, f"{request.dashboard_id}.json")
                
                if os.path.exists(dashboard_file):
                    with open(dashboard_file, 'r', encoding='utf-8') as f:
                        dashboard_data = json.load(f)
                    
                    # Add new conversation entries
                    conversation_history = dashboard_data.get('conversation_history', [])
                    
                    # Add user message
                    user_message = {
                        "id": len(conversation_history) + 1,
                        "type": "user",
                        "content": request.message,
                        "timestamp": datetime.now().isoformat()
                    }
                    conversation_history.append(user_message)
                    
                    # Add AI response
                    ai_message = {
                        "id": len(conversation_history) + 1,
                        "type": "ai",
                        "content": ai_response,
                        "sources": sources,
                        "timestamp": datetime.now().isoformat()
                    }
                    conversation_history.append(ai_message)
                    
                    # Update dashboard with new conversation history
                    dashboard_data['conversation_history'] = conversation_history
                    
                    # Save updated dashboard
                    with open(dashboard_file, 'w', encoding='utf-8') as f:
                        json.dump(dashboard_data, f, indent=2, ensure_ascii=False)
                        
                    print(f"Saved conversation history to dashboard {request.dashboard_id}")
                    
            except Exception as e:
                print(f"Error saving conversation history: {e}")
                # Don't fail the request if conversation saving fails
        
        return {
            "success": True,
            "response": ai_response,
            "sources": sources,
            "case_name": case_name,
            "timestamp": datetime.now().isoformat()
        }
        
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        print(f"Error in chat continuation: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing chat: {str(e)}")


@app.get("/api/cases/{case_name}/research-questions")
async def get_research_questions(case_name: str):
    """Get research questions for a case from ProjectLib."""
    try:
        project_file = os.path.join("ProjectLib", f"{case_name}.json")
        
        if not os.path.exists(project_file):
            raise HTTPException(status_code=404, detail=f"Project {case_name} not found")
        
        with open(project_file, 'r', encoding='utf-8') as f:
            project_data = json.load(f)
        
        research_questions = project_data.get('research_questions', [])
        
        return {
            "success": True,
            "case_name": case_name,
            "research_questions": research_questions
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error getting research questions: {e}")
        raise HTTPException(status_code=500, detail=f"Error loading research questions: {str(e)}")


@app.put("/api/cases/{case_name}/research-questions")
async def update_research_questions(case_name: str, request: dict):
    """Update research questions for a case in ProjectLib."""
    try:
        project_file = os.path.join("ProjectLib", f"{case_name}.json")
        
        if not os.path.exists(project_file):
            raise HTTPException(status_code=404, detail=f"Project {case_name} not found")
        
        # Load existing project data
        with open(project_file, 'r', encoding='utf-8') as f:
            project_data = json.load(f)
        
        # Update research questions
        research_questions = request.get('research_questions', [])
        project_data['research_questions'] = research_questions
        
        # Save updated project data
        with open(project_file, 'w', encoding='utf-8') as f:
            json.dump(project_data, f, indent=4, ensure_ascii=False)
        
        print(f"Updated research questions for {case_name}: {len(research_questions)} questions")
        
        return {
            "success": True,
            "case_name": case_name,
            "research_questions": research_questions,
            "message": f"Updated {len(research_questions)} research questions"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error updating research questions: {e}")
        raise HTTPException(status_code=500, detail=f"Error saving research questions: {str(e)}")